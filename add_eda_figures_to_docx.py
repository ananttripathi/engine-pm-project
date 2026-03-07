#!/usr/bin/env python3
"""Add all EDA figure images to the Interim Report docx in the correct sections."""
import os
from docx import Document
from docx.shared import Inches

DOCX_PATH = "AnantTripathi_EnginePredictiveMaintenance_InterimReport.docx"
FIG_DIR = "figures"

# (caption, filename) in display order
UNIVARIATE_FIGS = [
    ("Figure 1: Target distribution (Engine_Condition)", "fig1_target_distribution.png"),
    ("Figure 2: Univariate feature distributions (histograms)", "fig2_univariate_histograms.png"),
]
BIVARIATE_FIGS = [
    ("Figure 3: Box plots by Engine Condition", "fig3_boxplots.png"),
    ("Figure 4: Violin plots (subset of features)", "fig4_violin_subset.png"),
    ("Figure 5: Violin plots (all 6 features)", "fig5_violin_all6.png"),
    ("Figure 6: Strip plots (Lub_Oil_Temperature, Coolant_Temperature)", "fig6_strip_plots.png"),
    ("Figure 7: Mean feature values by Engine Condition", "fig7_mean_bar_chart.png"),
]
MULTIVARIATE_FIGS = [
    ("Figure 8: Correlation matrix", "fig8_correlation_heatmap.png"),
    ("Figure 9: Scatter plots (temperatures and RPM vs pressure)", "fig9_scatter_two.png"),
    ("Figure 10: Scatter matrix (feature pairs by condition)", "fig10_scatter_2x2.png"),
    ("Figure 11: Feature–target correlation (absolute)", "fig11_corr_bar.png"),
    ("Figure 12: Pair plot (sample, 3 features by condition)", "fig12_pairplot.png"),
]

def insert_figure_block(doc, body, insert_index, caption_text, image_path, width_inches=5.2):
    """Append a caption paragraph and an image paragraph to doc, then move them to body at insert_index."""
    if not os.path.exists(image_path):
        print("Skip (not found):", image_path)
        return 0
    # Add caption and image at end of document
    p_cap = doc.add_paragraph()
    r_cap = p_cap.add_run(caption_text)
    r_cap.italic = True
    r_cap.font.size = None  # keep default
    p_img = doc.add_paragraph()
    p_img.add_run().add_picture(image_path, width=Inches(width_inches))
    # Move to correct position: caption at insert_index, image at insert_index+1
    body_el = list(body)
    cap_elem = body_el[-2]
    img_elem = body_el[-1]
    body.remove(img_elem)
    body.insert(insert_index + 1, img_elem)
    body.remove(cap_elem)
    body.insert(insert_index, cap_elem)
    return 2

def main():
    doc = Document(DOCX_PATH)
    body = doc.element.body
    body_list = list(body)

    # Find body indices: insert figures AFTER the "Refer to notebook" line for each subsection
    # so figures appear at the end of that subsection's text.
    idx_ref_univariate = None   # "Refer to notebook: target bar chart"
    idx_ref_bivariate = None    # "Refer to notebook: box plots"
    idx_ref_multivariate = None # "Refer to notebook: correlation heatmap"
    for i, child in enumerate(body_list):
        if child.tag.endswith("p"):
            text = "".join(t for t in child.itertext())
            if "Refer to notebook" in text and "target bar chart" in text:
                idx_ref_univariate = i
            if "Refer to notebook" in text and "box plots" in text:
                idx_ref_bivariate = i
            if "Refer to notebook" in text and "correlation heatmap" in text:
                idx_ref_multivariate = i

    if idx_ref_univariate is None:
        for p in doc.paragraphs:
            if "Refer to notebook" in p.text and "histograms" in p.text:
                try:
                    idx_ref_univariate = body_list.index(p._element)
                except ValueError:
                    pass
                break
    if idx_ref_bivariate is None:
        for p in doc.paragraphs:
            if "Refer to notebook" in p.text and "box plots" in p.text:
                try:
                    idx_ref_bivariate = body_list.index(p._element)
                except ValueError:
                    pass
                break
    if idx_ref_multivariate is None:
        for p in doc.paragraphs:
            if "Refer to notebook" in p.text and "correlation" in p.text:
                try:
                    idx_ref_multivariate = body_list.index(p._element)
                except ValueError:
                    pass
                break

    # Insert right after the "Refer to notebook" line for each subsection
    pos_univariate = idx_ref_univariate + 1 if idx_ref_univariate is not None else 46
    pos_bivariate = idx_ref_bivariate + 1 if idx_ref_bivariate is not None else 58
    pos_multivariate = idx_ref_multivariate + 1 if idx_ref_multivariate is not None else 62

    # Insert univariate first, then bivariate, then multivariate (so later positions stay valid)
    for caption, fname in reversed(UNIVARIATE_FIGS):
        path = os.path.join(FIG_DIR, fname)
        n = insert_figure_block(doc, body, pos_univariate, caption, path)
        pos_univariate += n
    pos_bivariate += 4   # we added 4 elements (2 univariate figures)
    pos_multivariate += 4

    for caption, fname in reversed(BIVARIATE_FIGS):
        path = os.path.join(FIG_DIR, fname)
        n = insert_figure_block(doc, body, pos_bivariate, caption, path)
        pos_bivariate += n
    pos_multivariate += 10  # 5 bivariate figures * 2

    for caption, fname in reversed(MULTIVARIATE_FIGS):
        path = os.path.join(FIG_DIR, fname)
        n = insert_figure_block(doc, body, pos_multivariate, caption, path)
        pos_multivariate += n

    doc.save(DOCX_PATH)
    print("Saved:", DOCX_PATH)
    print("EDA figures added to report.")

if __name__ == "__main__":
    main()
