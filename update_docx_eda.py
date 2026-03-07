#!/usr/bin/env python3
"""Update the EDA section (Section 3) in the Interim Report docx."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_PATH = "AnantTripathi_EnginePredictiveMaintenance_InterimReport.docx"
EDA_START_IDX = 32   # body child index for "3. Exploratory Data Analysis"
EDA_END_IDX = 89     # body child index for "4. Data Preparation" (exclusive: we remove 32..88)

def main():
    doc = Document(DOCX_PATH)
    body = doc.element.body
    children = list(body)

    # Remove existing EDA section (paragraphs and tables)
    for i in range(EDA_END_IDX - 1, EDA_START_IDX - 1, -1):
        body.remove(children[i])

    # Build new EDA content (each add_paragraph/add_table appends to body)
    new_paras = [
        ("3. Exploratory Data Analysis", True),   # bold for heading
        ("3.1 Data Collection and Background", True),
        ("The dataset consists of engine sensor readings collected from equipment in operation. "
         "Each row represents a snapshot of key physical measurements—RPM, oil pressure, fuel pressure, "
         "coolant pressure, oil and coolant temperatures—along with an Engine_Condition label (Normal vs. "
         "Maintenance Required). This data supports predictive maintenance: by learning patterns that precede "
         "failures, we can build models to flag engines before they break down. The same dataset is registered "
         "on Hugging Face for reproducibility and version control.", False),
        ("3.2 Data Overview", True),
    ]
    for text, bold in new_paras:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold

    # Overview as paragraphs (table can be re-added in Word if desired)
    doc.add_paragraph("Rows: 19,535  |  Columns: 7 (6 features + 1 target)  |  Missing Values: None  |  Data Types: All numerical (int64: Engine_RPM, target; float64: pressures, temperatures)")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Features: ").bold = True
    p.add_run("Engine_RPM, Lub_Oil_Pressure, Fuel_Pressure, Coolant_Pressure, Lub_Oil_Temperature, Coolant_Temperature  ")
    p = doc.add_paragraph()
    p.add_run("Target: ").bold = True
    p.add_run("Engine_Condition (0: Normal, 1: Maintenance Required)")
    doc.add_paragraph("Basic statistics (min, max, mean, std, quartiles) confirm ranges consistent with engine sensor data. "
                      "The target is imbalanced (e.g. ~63% Maintenance, ~37% Normal), which motivates stratified splits and F1-based evaluation.")

    more_paras = [
        ("3.3 Univariate Analysis", True),
        ("Target distribution: A bar chart of Engine_Condition shows class counts (Normal vs. Maintenance). The imbalance supports using stratified sampling and metrics such as F1.", False),
        ("Feature distributions: Histograms and summary statistics (mean, std, skewness) describe each numeric feature. Ranges and central tendency are documented; some features show skew or long tails (e.g. Coolant_Temperature with high max and positive skew). For tree-based models we proceed with raw features; scaling is applied in the pipeline for consistency.", False),
        ("Observation: Class imbalance and feature skewness inform preprocessing and model choice (stratified splits, F1, and tree-based algorithms that handle mixed scales).", False),
        ("Refer to notebook: target bar chart, histograms per feature, and feature statistics.", False),
        ("3.4 Bivariate Analysis", True),
        ("Bivariate analysis compares each feature across the two target classes (Normal vs. Maintenance) to identify which sensors are most discriminative.", False),
        ("Box plots: Features plotted by Engine_Condition show differences in median, spread, and outliers between classes. All six sensor features are compared side by side.", False),
        ("Violin plots: First a subset of four features, then all six features in a 2×3 grid. Violins show full distribution shape and density by class, making overlap and separation visible (e.g. temperatures and pressures often differ between Normal and Maintenance).", False),
        ("Strip plots: Point clouds for Lub_Oil_Temperature and Coolant_Temperature by condition (on a sample) illustrate overlap and spread of individual readings.", False),
        ("Grouped bar chart: Mean of each feature by Engine_Condition gives a direct numeric comparison; the chart highlights which sensors differ most between classes.", False),
        ("Mean table: Mean feature values by class are printed for exact numeric comparison.", False),
        ("Business insight: Features that separate the two classes (e.g. pressures, temperatures, RPM) are strong candidates for the model. Real-time monitoring of these sensors can support proactive maintenance scheduling.", False),
        ("Refer to notebook: box plots, violin plots (subset and all 6), strip plots, grouped bar chart, and mean-by-class table.", False),
        ("3.5 Multivariate Analysis", True),
        ("Multivariate analysis examines correlations between features and with the target, and how multiple variables relate together.", False),
        ("Correlation matrix: A heatmap of pairwise correlations (including the target) shows linear relationships. Strong correlations between some features may suggest redundancy; tree-based models can still use multiple signals effectively.", False),
        ("Scatter plots by condition: Lub_Oil_Temperature vs. Coolant_Temperature and Engine_RPM vs. Lub_Oil_Pressure (colored by Engine_Condition) show how two dimensions separate the classes. An additional 2×2 grid of scatter plots covers: Fuel_Pressure vs. Coolant_Pressure; Coolant_Temperature vs. Coolant_Pressure; Engine_RPM vs. Coolant_Temperature; Lub_Oil_Pressure vs. Lub_Oil_Temperature. Together these reveal which feature pairs help distinguish Normal from Maintenance.", False),
        ("Feature–target correlation: A table and a horizontal bar chart of the absolute correlation of each feature with Engine_Condition quantify which features are most linearly associated with the target. Engine_RPM typically shows the strongest (negative) correlation with maintenance need.", False),
        ("Pair plot: A pair plot on a sample (e.g. Engine_RPM, Lub_Oil_Temperature, Coolant_Temperature) with KDE on the diagonal and scatter off-diagonal, colored by condition, summarizes joint structure and separation.", False),
        ("Observation: Features with higher correlation to the target and clearer separation in scatter plots support the feasibility of a classifier. Lower RPM combined with abnormal pressure/temperature patterns may indicate fault conditions.", False),
        ("Refer to notebook: correlation heatmap, scatter pairs (initial two + 2×2 grid), feature–target correlation bar chart, and pair plot.", False),
        ("3.6 EDA Insights and Observations", True),
        ("1. Data quality: All features are numeric; no missing values. Duplicates are dropped in data preparation.", False),
        ("2. Target balance: Engine_Condition is imbalanced (e.g. Maintenance ~63%, Normal ~37%). Stratified train/test splits and F1-score are used for tuning and evaluation.", False),
        ("3. Univariate: Feature distributions show different ranges and skewness; the target bar chart confirms imbalance.", False),
        ("4. Bivariate: Box and violin plots by condition show clear separation for several sensors; mean values by class and strip/bar visuals identify the most predictive features.", False),
        ("5. Multivariate: Correlation matrix and scatter plots show relationships between features and with the target; the feature–target correlation bar chart and pair plot support model and feature choices.", False),
        ("6. Recommendations applied: Stratified train/test split; all six sensor features retained; tree-based models (Random Forest, AdaBoost, Gradient Boosting); F1 and ROC-AUC monitored.", False),
    ]
    for text, bold in more_paras:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold

    # Move all newly added content (last N elements) to position EDA_START_IDX, preserving order.
    # New content is at end of body; insert last element first at highest index so first-added ends at EDA_START_IDX.
    body_children = list(body)
    n_new = len(body_children) - (len(children) - (EDA_END_IDX - EDA_START_IDX))
    start_old = len(body_children) - n_new
    for i in range(n_new - 1, -1, -1):  # i = n_new-1 down to 0
        elem = body_children[start_old + i]
        body.remove(elem)
        body.insert(EDA_START_IDX + i, elem)

    doc.save(DOCX_PATH)
    print("Saved:", DOCX_PATH)
    print("EDA section updated successfully.")

if __name__ == "__main__":
    main()
